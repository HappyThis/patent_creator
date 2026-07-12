from __future__ import annotations

import json
import re
import shutil
import subprocess
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable, cast
from uuid import uuid4

from docx import Document as create_document
from docx.document import Document as DocxDocument
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.styles.style import ParagraphStyle

from .disclosure import build_render_ast
from .figures import FIGURE_HEIGHT, FIGURE_WIDTH, figure_caption

FIGURE_REF_PATTERN = re.compile(r"\[([^\]]+)\]\(figure:(fig_\d{6})\)")
FORMULA_REF_PATTERN = re.compile(r"\[([^\]]+)\]\(formula:([A-Za-z0-9_-]+)\)")

BODY_FONT_PT = 12
BODY_LINE_SPACING = 1.9
BODY_PARAGRAPH_AFTER_PT = 9
LATIN_FONT = "Times New Roman"
EAST_ASIA_FONT = "宋体"
HEADING_1_PT = 18
HEADING_2_PT = 14.25
HEADING_BEFORE_PT = 25.5
HEADING_AFTER_PT = 12
FIGURE_BEFORE_PT = 21
FIGURE_CAPTION_AFTER_PT = 27
FORMULA_BEFORE_PT = 13.5
FORMULA_AFTER_PT = 16.5
TABLE_BEFORE_PT = 15
TABLE_AFTER_PT = 21
SPACER_BASE_LINE_PT = 12
FORMULA_LEFT_COL_IN = 0.25
FORMULA_BODY_COL_IN = 6.05
FORMULA_NUMBER_COL_IN = 0.5
FORMULA_MAX_WIDTH_IN = 5.9


@dataclass(frozen=True)
class InlineToken:
    type: str
    start: int
    end: int
    text: str = ""
    label: str = ""
    target_id: str = ""


class DocxExportError(RuntimeError):
    """Raised when DOCX asset rendering or export fails."""


def referenced_figure_block_ids(disclosure: dict[str, Any]) -> set[str]:
    """Return figure ids whose images are embedded as figure blocks."""

    figure_ids: set[str] = set()

    def visit_sections(sections: Iterable[dict[str, Any]]) -> None:
        for section in sections:
            for block in section.get("blocks", []):
                if not isinstance(block, dict) or block.get("type") != "figure":
                    continue
                figure_id = str(block.get("figure_id") or "")
                if figure_id:
                    figure_ids.add(figure_id)
            visit_sections(section.get("sections", []))

    visit_sections(disclosure.get("sections", []))
    return figure_ids


def export_disclosure_docx(
    *,
    disclosure: dict[str, Any],
    figures: list[dict[str, Any]],
    export_path: Path,
    project_dir: Path,
) -> Path:
    render_ast = build_render_ast(disclosure, figures=figures)
    figures_by_id = {str(figure.get("figure_id")): figure for figure in figures}
    formula_numbers = collect_formula_numbers(render_ast.get("children", []))
    asset_manifest, asset_dir = render_assets(render_ast, figures_by_id, project_dir)

    try:
        document = create_document()
        configure_document(document)

        for index, node in enumerate(render_ast.get("children", []), start=1):
            if node.get("type") == "section":
                render_section(
                    document,
                    node,
                    index_path=[index],
                    figures_by_id=figures_by_id,
                    formula_numbers=formula_numbers,
                    assets=asset_manifest,
                )

        export_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(str(export_path))
    finally:
        if asset_dir is not None:
            shutil.rmtree(asset_dir, ignore_errors=True)
    return export_path


def configure_document(document: DocxDocument) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)

    styles = document.styles
    normal = get_paragraph_style(styles, "Normal")
    set_style_font(normal)
    normal.font.size = Pt(BODY_FONT_PT)
    normal.font.color.rgb = RGBColor(0x2E, 0x39, 0x42)
    normal.paragraph_format.line_spacing = BODY_LINE_SPACING
    normal.paragraph_format.space_after = Pt(BODY_PARAGRAPH_AFTER_PT)

    for name in ("Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"):
        style = get_paragraph_style(styles, name)
        set_style_font(style)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0x0D, 0x16, 0x24)
        style.paragraph_format.space_before = Pt(HEADING_BEFORE_PT)
        style.paragraph_format.space_after = Pt(HEADING_AFTER_PT)
        style.paragraph_format.line_spacing = 1.35
    get_paragraph_style(styles, "Heading 1").font.size = Pt(HEADING_1_PT)
    get_paragraph_style(styles, "Heading 2").font.size = Pt(HEADING_2_PT)
    get_paragraph_style(styles, "Heading 3").font.size = Pt(HEADING_2_PT)
    for name in ("List Bullet", "List Number"):
        style = get_paragraph_style(styles, name)
        style.font.bold = False
        style.font.size = Pt(BODY_FONT_PT)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = BODY_LINE_SPACING


def get_paragraph_style(styles: Any, name: str) -> ParagraphStyle:
    return cast(ParagraphStyle, styles[name])


def set_style_font(style: ParagraphStyle) -> None:
    style.font.name = LATIN_FONT
    style_element = style._element
    if style_element is None:
        return
    r_pr = style_element.get_or_add_rPr()
    set_r_fonts(r_pr)


def set_run_font(run: Any) -> None:
    run.font.name = LATIN_FONT
    r_pr = run._element.get_or_add_rPr()
    set_r_fonts(r_pr)


def set_r_fonts(r_pr: Any) -> None:
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), LATIN_FONT)
    r_fonts.set(qn("w:hAnsi"), LATIN_FONT)
    r_fonts.set(qn("w:cs"), LATIN_FONT)
    r_fonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)


def collect_formula_numbers(nodes: Iterable[dict[str, Any]]) -> dict[str, int]:
    formula_numbers: dict[str, int] = {}

    def visit(items: Iterable[dict[str, Any]]) -> None:
        for item in items:
            if item.get("type") == "section":
                visit(item.get("children", []))
            elif item.get("type") == "formula":
                formula_numbers[str(item["id"])] = len(formula_numbers) + 1

    visit(nodes)
    return formula_numbers


def render_assets(
    render_ast: dict[str, Any],
    figures_by_id: dict[str, dict[str, Any]],
    project_dir: Path,
) -> tuple[dict[str, dict[str, Any]], Path | None]:
    items: list[dict[str, str]] = []
    existing_figure_assets: dict[str, dict[str, Any]] = {}

    def add_inline_math(latex: str) -> str:
        asset_id = f"inline_{len(items) + 1:06d}_{uuid4().hex[:8]}"
        items.append({"id": asset_id, "kind": "inline_formula", "latex": latex})
        return asset_id

    inline_asset_by_latex: dict[str, str] = {}
    block_formula_ids: set[str] = set()
    figure_ids: set[str] = set()

    def scan_text(text: str) -> None:
        for token in parse_inline_tokens(text):
            if token.type == "math" and token.text not in inline_asset_by_latex:
                inline_asset_by_latex[token.text] = add_inline_math(token.text)

    def visit(nodes: Iterable[dict[str, Any]]) -> None:
        for node in nodes:
            node_type = node.get("type")
            if node_type == "section":
                visit(node.get("children", []))
            elif node_type in {"title", "paragraph"}:
                scan_text(str(node.get("text") or ""))
            elif node_type == "list":
                for item in node.get("items", []):
                    scan_text(str(item))
            elif node_type == "table":
                for value in node.get("columns", []):
                    scan_text(str(value))
                for row in node.get("rows", []):
                    for value in row:
                        scan_text(str(value))
            elif node_type == "formula":
                block_id = str(node["id"])
                if block_id not in block_formula_ids:
                    block_formula_ids.add(block_id)
                    items.append({"id": f"block_{block_id}", "kind": "block_formula", "latex": str(node.get("latex") or "")})
            elif node_type == "figure":
                figure_id = str(node.get("figure_id") or "")
                figure = figures_by_id.get(figure_id)
                if figure_id and isinstance(figure, dict) and figure_id not in figure_ids:
                    figure_ids.add(figure_id)
                    asset = _existing_figure_asset(figure_id, figure, project_dir)
                    if asset:
                        existing_figure_assets[asset["id"]] = asset

    visit(render_ast.get("children", []))
    if not items:
        return existing_figure_assets, None

    asset_dir = project_dir / "exports" / f"docx-assets-{uuid4().hex[:8]}"
    input_path = asset_dir / "input.json"
    output_dir = asset_dir / "images"
    manifest_path = asset_dir / "manifest.json"
    asset_dir.mkdir(parents=True, exist_ok=True)
    try:
        input_path.write_text(json.dumps({"items": items}, ensure_ascii=False, indent=2), encoding="utf-8")

        repo_root = Path(__file__).resolve().parents[3]
        frontend_root = repo_root / "frontend"
        script_path = frontend_root / "scripts" / "render-docx-assets.mjs"
        result = subprocess.run(
            [
                "node",
                str(script_path),
                "--input",
                str(input_path),
                "--output",
                str(output_dir),
                "--manifest",
                str(manifest_path),
            ],
            cwd=frontend_root,
            capture_output=True,
            text=True,
            check=False,
            timeout=90,
        )
        if result.returncode != 0:
            message = result.stderr.strip() or result.stdout.strip() or "unknown renderer error"
            raise DocxExportError(f"DOCX asset renderer failed: {message}")

        manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
        assets = {**existing_figure_assets, **_validate_asset_manifest(manifest, items, output_dir)}
    except OSError as exc:
        shutil.rmtree(asset_dir, ignore_errors=True)
        raise DocxExportError(f"DOCX asset renderer failed to start: {exc}") from exc
    except subprocess.TimeoutExpired as exc:
        shutil.rmtree(asset_dir, ignore_errors=True)
        raise DocxExportError("DOCX asset renderer timed out.") from exc
    except json.JSONDecodeError as exc:
        shutil.rmtree(asset_dir, ignore_errors=True)
        raise DocxExportError(f"DOCX asset renderer returned invalid JSON: {exc}") from exc
    except DocxExportError:
        shutil.rmtree(asset_dir, ignore_errors=True)
        raise

    for latex, asset_id in inline_asset_by_latex.items():
        if asset_id in assets:
            assets[f"inline:{latex}"] = assets[asset_id]
    return assets, asset_dir


def _existing_figure_asset(figure_id: str, figure: dict[str, Any], project_dir: Path) -> dict[str, Any] | None:
    render = figure.get("render", {})
    if not isinstance(render, dict) or render.get("type") != "png":
        return None
    render_path = render.get("path")
    if not isinstance(render_path, str) or not render_path:
        raise DocxExportError(f"DOCX figure render path is missing: {figure_id}")
    project_root = project_dir.resolve()
    resolved_path = (project_root / render_path).resolve()
    if not resolved_path.is_relative_to(project_root):
        raise DocxExportError(f"DOCX figure render path is outside the project: {figure_id}")
    if not resolved_path.is_file():
        raise DocxExportError(f"DOCX figure render file is missing: {figure_id}")
    width_px = render.get("width") if isinstance(render.get("width"), (int, float)) else FIGURE_WIDTH
    height_px = render.get("height") if isinstance(render.get("height"), (int, float)) else FIGURE_HEIGHT
    return {
        "id": f"figure_{figure_id}",
        "kind": "figure",
        "path": str(resolved_path),
        "width_px": width_px,
        "height_px": height_px,
    }


def _validate_asset_manifest(
    manifest: Any,
    items: list[dict[str, str]],
    output_dir: Path,
) -> dict[str, dict[str, Any]]:
    if not isinstance(manifest, dict) or not isinstance(manifest.get("assets"), dict):
        raise DocxExportError("DOCX asset renderer returned an invalid manifest.")

    assets: dict[str, dict[str, Any]] = {}
    raw_assets = manifest["assets"]
    output_root = output_dir.resolve()
    for item in items:
        asset_id = item["id"]
        raw_asset = raw_assets.get(asset_id)
        if not isinstance(raw_asset, dict):
            raise DocxExportError(f"DOCX asset renderer did not return asset: {asset_id}")
        asset_path = raw_asset.get("path")
        width_px = raw_asset.get("width_px")
        height_px = raw_asset.get("height_px")
        if not isinstance(asset_path, str) or not asset_path:
            raise DocxExportError(f"DOCX asset renderer returned an invalid path for asset: {asset_id}")
        resolved_asset_path = Path(asset_path).resolve()
        if not resolved_asset_path.is_relative_to(output_root):
            raise DocxExportError(f"DOCX asset renderer returned an out-of-directory path for asset: {asset_id}")
        if not resolved_asset_path.is_file():
            raise DocxExportError(f"DOCX asset renderer returned a missing file for asset: {asset_id}")
        if not isinstance(width_px, (int, float)) or width_px <= 0:
            raise DocxExportError(f"DOCX asset renderer returned an invalid width for asset: {asset_id}")
        if not isinstance(height_px, (int, float)) or height_px <= 0:
            raise DocxExportError(f"DOCX asset renderer returned an invalid height for asset: {asset_id}")
        assets[asset_id] = {**raw_asset, "path": str(resolved_asset_path)}
    return assets


def render_section(
    document: DocxDocument,
    node: dict[str, Any],
    *,
    index_path: list[int],
    figures_by_id: dict[str, dict[str, Any]],
    formula_numbers: dict[str, int],
    assets: dict[str, dict[str, Any]],
) -> None:
    heading = document.add_paragraph()
    configure_heading_paragraph(heading, level=len(index_path), is_first=len(index_path) == 1 and index_path[0] == 1)
    prefix = heading.add_run(f"{'.'.join(str(part) for part in index_path)}. ")
    style_heading_run(prefix, level=len(index_path))
    title = heading.add_run(str(node.get("title") or "").strip())
    style_heading_run(title, level=len(index_path))

    children = node.get("children", [])
    has_child_sections = any(child.get("type") == "section" for child in children)
    has_direct_content = any(child.get("type") != "section" for child in children)
    if not has_child_sections and not has_direct_content:
        return

    child_section_index = 0
    for child in children:
        if child.get("type") == "section":
            child_section_index += 1
            render_section(
                document,
                child,
                index_path=[*index_path, child_section_index],
                figures_by_id=figures_by_id,
                formula_numbers=formula_numbers,
                assets=assets,
            )
        else:
            render_block(document, child, figures_by_id=figures_by_id, formula_numbers=formula_numbers, assets=assets)


def render_block(
    document: DocxDocument,
    node: dict[str, Any],
    *,
    figures_by_id: dict[str, dict[str, Any]],
    formula_numbers: dict[str, int],
    assets: dict[str, dict[str, Any]],
) -> None:
    node_type = node.get("type")
    if node_type in {"title", "paragraph"}:
        paragraph = document.add_paragraph()
        configure_body_paragraph(paragraph)
        if node_type == "paragraph":
            paragraph.paragraph_format.first_line_indent = Pt(BODY_FONT_PT * 2)
        append_inline_content(paragraph, str(node.get("text") or ""), figures_by_id, formula_numbers, assets)
    elif node_type == "list":
        style = "List Number" if node.get("ordered") else "List Bullet"
        for item in node.get("items", []):
            paragraph = document.add_paragraph(style=style)
            configure_body_paragraph(paragraph, after_pt=6)
            append_inline_content(paragraph, str(item), figures_by_id, formula_numbers, assets)
    elif node_type == "table":
        render_table(document, node, figures_by_id, formula_numbers, assets)
    elif node_type == "formula":
        render_formula(document, node, formula_numbers, assets)
    elif node_type == "figure":
        render_figure(document, node, figures_by_id, assets)
    elif node_type == "image":
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(str(node.get("alt") or node.get("src") or "image"))
        set_run_font(run)
        run.italic = True
        run.font.color.rgb = RGBColor(0x6F, 0x77, 0x80)


def render_table(
    document: DocxDocument,
    node: dict[str, Any],
    figures_by_id: dict[str, dict[str, Any]],
    formula_numbers: dict[str, int],
    assets: dict[str, dict[str, Any]],
) -> None:
    columns = [str(value) for value in node.get("columns", [])]
    rows = [[str(value) for value in row] for row in node.get("rows", [])]
    if not columns:
        return

    add_vertical_space(document, TABLE_BEFORE_PT)
    table = document.add_table(rows=1, cols=len(columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table, "D4DBE6")

    for index, column in enumerate(columns):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, "E8E0D4")
        set_cell_margins(cell, top=150, start=180, bottom=150, end=180)
        paragraph = cell.paragraphs[0]
        configure_table_paragraph(paragraph)
        append_inline_content(paragraph, column, figures_by_id, formula_numbers, assets)
        for run in paragraph.runs:
            run.bold = True

    for row in rows:
        cells = table.add_row().cells
        for index, cell in enumerate(cells):
            set_cell_margins(cell, top=150, start=180, bottom=150, end=180)
            paragraph = cell.paragraphs[0]
            configure_table_paragraph(paragraph)
            append_inline_content(paragraph, row[index] if index < len(row) else "", figures_by_id, formula_numbers, assets)

    add_vertical_space(document, TABLE_AFTER_PT)


def render_formula(
    document: DocxDocument,
    node: dict[str, Any],
    formula_numbers: dict[str, int],
    assets: dict[str, dict[str, Any]],
) -> None:
    asset = assets.get(f"block_{node.get('id')}")
    number = formula_numbers.get(str(node.get("id")))
    add_vertical_space(document, FORMULA_BEFORE_PT)
    table = document.add_table(rows=1, cols=3)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(table)
    table.columns[0].width = Inches(FORMULA_LEFT_COL_IN)
    table.columns[1].width = Inches(FORMULA_BODY_COL_IN)
    table.columns[2].width = Inches(FORMULA_NUMBER_COL_IN)
    left_cell, formula_cell, number_cell = table.rows[0].cells
    left_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    formula_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    number_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_width(left_cell, FORMULA_LEFT_COL_IN)
    set_cell_width(formula_cell, FORMULA_BODY_COL_IN)
    set_cell_width(number_cell, FORMULA_NUMBER_COL_IN)
    set_cell_margins(left_cell, top=0, start=0, bottom=0, end=0)
    set_cell_margins(formula_cell, top=0, start=0, bottom=0, end=0)
    set_cell_margins(number_cell, top=0, start=0, bottom=0, end=0)

    paragraph = formula_cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing = 1.0
    if asset:
        add_picture(paragraph, asset, max_width=Inches(FORMULA_MAX_WIDTH_IN))
    else:
        append_plain_text(paragraph, str(node.get("latex") or ""))

    number_paragraph = number_cell.paragraphs[0]
    number_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    number_paragraph.paragraph_format.line_spacing = 1.0
    if number:
        run = number_paragraph.add_run(f"({number})")
        set_run_font(run)
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0x6F, 0x77, 0x80)
    add_vertical_space(document, FORMULA_AFTER_PT)


def render_figure(
    document: DocxDocument,
    node: dict[str, Any],
    figures_by_id: dict[str, dict[str, Any]],
    assets: dict[str, dict[str, Any]],
) -> None:
    figure_id = str(node.get("figure_id") or "")
    figure = figures_by_id.get(figure_id)
    asset = assets.get(f"figure_{figure_id}")
    add_vertical_space(document, FIGURE_BEFORE_PT)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing = 1.0
    if asset:
        add_picture(paragraph, asset, max_width=Inches(5.7))
    else:
        run = paragraph.add_run(figure_id)
        set_run_font(run)
        run.italic = True
        run.font.color.rgb = RGBColor(0x6F, 0x77, 0x80)

    if figure:
        caption = document.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.space_before = Pt(7.5)
        caption.paragraph_format.space_after = Pt(FIGURE_CAPTION_AFTER_PT)
        caption.paragraph_format.line_spacing = 1.55
        run = caption.add_run(figure_caption(figure))
        set_run_font(run)
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x5F, 0x66, 0x70)


def append_inline_content(
    paragraph: Any,
    text: str,
    figures_by_id: dict[str, dict[str, Any]],
    formula_numbers: dict[str, int],
    assets: dict[str, dict[str, Any]],
) -> None:
    cursor = 0
    for token in parse_inline_tokens(text):
        if token.start > cursor:
            append_plain_text(paragraph, unescape_inline_text(text[cursor:token.start]))
        if token.type == "math":
            asset = assets.get(f"inline:{token.text}")
            if asset:
                add_picture(paragraph, asset, inline=True)
            else:
                append_plain_text(paragraph, token.label)
        elif token.type == "figure":
            figure = figures_by_id.get(token.target_id)
            run = paragraph.add_run(str(figure.get("label") if figure else token.label))
            set_run_font(run)
            style_reference_run(run)
        elif token.type == "formula":
            number = formula_numbers.get(token.target_id)
            run = paragraph.add_run(f"式({number})" if number else token.label)
            set_run_font(run)
            style_reference_run(run)
        cursor = token.end
    if cursor < len(text):
        append_plain_text(paragraph, unescape_inline_text(text[cursor:]))


def append_plain_text(paragraph: Any, text: str) -> None:
    if text:
        run = paragraph.add_run(text)
        set_run_font(run)


def parse_inline_tokens(text: str) -> list[InlineToken]:
    tokens: list[InlineToken] = []
    cursor = 0
    while cursor < len(text):
        candidates = [
            find_next_figure_ref(text, cursor),
            find_next_formula_ref(text, cursor),
            find_next_inline_math(text, cursor),
        ]
        token = min((candidate for candidate in candidates if candidate), key=lambda item: item.start, default=None)
        if token is None:
            break
        tokens.append(token)
        cursor = token.end
    return tokens


def find_next_figure_ref(text: str, offset: int) -> InlineToken | None:
    match = FIGURE_REF_PATTERN.search(text, offset)
    if not match:
        return None
    return InlineToken("figure", match.start(), match.end(), label=match.group(1), target_id=match.group(2))


def find_next_formula_ref(text: str, offset: int) -> InlineToken | None:
    match = FORMULA_REF_PATTERN.search(text, offset)
    if not match:
        return None
    return InlineToken("formula", match.start(), match.end(), label=match.group(1), target_id=match.group(2))


def find_next_inline_math(text: str, offset: int) -> InlineToken | None:
    for start in range(offset, len(text)):
        if text[start] != "$" or is_escaped(text, start) or (start + 1 < len(text) and text[start + 1] == "$"):
            continue
        for end in range(start + 1, len(text)):
            if text[end] == "\n":
                break
            if text[end] == "$" and not is_escaped(text, end):
                latex = text[start + 1 : end].strip()
                if not latex:
                    break
                return InlineToken("math", start, end + 1, text=latex, label=text[start : end + 1])
    return None


def is_escaped(text: str, index: int) -> bool:
    slash_count = 0
    cursor = index - 1
    while cursor >= 0 and text[cursor] == "\\":
        slash_count += 1
        cursor -= 1
    return slash_count % 2 == 1


def unescape_inline_text(text: str) -> str:
    return text.replace(r"\$", "$")


def style_reference_run(run: Any) -> None:
    run.font.color.rgb = RGBColor(0x8A, 0x64, 0x26)
    run.underline = True


def configure_heading_paragraph(paragraph: Any, *, level: int, is_first: bool = False) -> None:
    paragraph.paragraph_format.space_before = Pt(0 if is_first else HEADING_BEFORE_PT)
    paragraph.paragraph_format.space_after = Pt(HEADING_AFTER_PT)
    paragraph.paragraph_format.line_spacing = 1.35
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.keep_together = True


def style_heading_run(run: Any, *, level: int) -> None:
    set_run_font(run)
    run.font.bold = True
    run.font.size = Pt(HEADING_1_PT if level == 1 else HEADING_2_PT)
    run.font.color.rgb = RGBColor(0x0D, 0x16, 0x24)


def configure_body_paragraph(paragraph: Any, *, after_pt: float = BODY_PARAGRAPH_AFTER_PT) -> None:
    paragraph.paragraph_format.line_spacing = BODY_LINE_SPACING
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after_pt)


def configure_table_paragraph(paragraph: Any) -> None:
    paragraph.paragraph_format.line_spacing = BODY_LINE_SPACING
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)


def add_vertical_space(document: DocxDocument, height_pt: float) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(max(height_pt - SPACER_BASE_LINE_PT, 0))
    run = paragraph.add_run()
    run.font.size = Pt(1)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.add_text(" ")


def add_picture(paragraph: Any, asset: dict[str, Any], *, max_width: Any | None = None, inline: bool = False) -> None:
    path = str(asset.get("path") or "")
    if not path:
        return
    width_px = max(float(asset.get("width_px") or 1), 1.0)
    height_px = max(float(asset.get("height_px") or 1), 1.0)
    run = paragraph.add_run()
    if inline:
        height_pt = min(max(height_px * 0.42, 8.5), 12.5)
        run.add_picture(path, height=Pt(height_pt))
        return
    if max_width is None:
        run.add_picture(path)
        return
    max_width_inches = max_width.inches
    if width_px / height_px < 0.55:
        max_width_inches = min(max_width_inches, 2.1)
    target_width_inches = min(width_px / 96.0, max_width_inches)
    run.add_picture(path, width=Inches(target_width_inches))


def set_table_borders(table: Any, color: str) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def remove_table_borders(table: Any) -> None:
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        borders.append(element)
    tbl_pr.append(borders)


def set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell: Any, *, top: int, start: int, bottom: int, end: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_cell_width(cell: Any, width_inches: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    width = tc_pr.find(qn("w:tcW"))
    if width is None:
        width = OxmlElement("w:tcW")
        tc_pr.append(width)
    width.set(qn("w:w"), str(int(width_inches * 1440)))
    width.set(qn("w:type"), "dxa")
