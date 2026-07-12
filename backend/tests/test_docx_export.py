from __future__ import annotations

import json
import subprocess
import threading
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
import pytest

from app.domain.disclosure import build_initial_disclosure
from app.domain import docx_export
from app.domain.docx_export import DocxExportError, export_disclosure_docx
from app.storage import workspace_store as workspace_store_module
from app.storage.workspace_store import WorkspaceStore

PNG_BYTES = (
    b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
    b"\x08\x02\x00\x00\x00\x90wS\xde\x00\x00\x00\x0cIDATx\x9cc\xf8\xff\xff?"
    b"\x00\x05\xfe\x02\xfeA\xd7S\x84\x00\x00\x00\x00IEND\xaeB`\x82"
)


def test_docx_export_renders_headings_without_word_outline_markers(tmp_path: Path) -> None:
    disclosure = build_initial_disclosure("heading export")

    output_path = export_disclosure_docx(
        disclosure=disclosure,
        figures=[],
        export_path=tmp_path / "heading-export.docx",
        project_dir=tmp_path,
    )

    document = Document(str(output_path))
    heading = next(paragraph for paragraph in document.paragraphs if paragraph.text.startswith("1. "))
    assert heading.style is not None
    assert heading.style.name == "Normal"
    assert heading.runs[0].bold is True
    assert heading.runs[0].font.color.rgb == docx_export.RGBColor(0x0D, 0x16, 0x24)


def test_docx_export_uses_songti_for_chinese_and_times_for_latin(tmp_path: Path) -> None:
    disclosure = build_initial_disclosure("font export")
    disclosure["sections"][0]["blocks"] = [
        {
            "id": "blk_font_test",
            "type": "paragraph",
            "text": "中文 ABC 123",
        }
    ]

    output_path = export_disclosure_docx(
        disclosure=disclosure,
        figures=[],
        export_path=tmp_path / "font-export.docx",
        project_dir=tmp_path,
    )

    document = Document(str(output_path))
    paragraph = next(paragraph for paragraph in document.paragraphs if "中文 ABC 123" in paragraph.text)
    r_pr = paragraph.runs[0]._element.rPr
    assert r_pr is not None
    r_fonts = r_pr.rFonts
    assert r_fonts.get(qn("w:ascii")) == "Times New Roman"
    assert r_fonts.get(qn("w:hAnsi")) == "Times New Roman"
    assert r_fonts.get(qn("w:eastAsia")) == "宋体"


def test_docx_export_leaves_empty_sections_blank(tmp_path: Path) -> None:
    disclosure = build_initial_disclosure("empty export")

    output_path = export_disclosure_docx(
        disclosure=disclosure,
        figures=[],
        export_path=tmp_path / "empty-export.docx",
        project_dir=tmp_path,
    )

    document = Document(str(output_path))
    assert "内容待补充" not in [paragraph.text for paragraph in document.paragraphs]


def test_docx_export_centers_block_formulas_across_full_line(tmp_path: Path) -> None:
    disclosure = build_initial_disclosure("formula export")
    disclosure["sections"][0]["blocks"].append(
        {
            "id": "blk_formula_test",
            "type": "formula",
            "latex": r"U=\sum_{i=1}^{n}\alpha_i L_i+\beta G",
        }
    )

    output_path = export_disclosure_docx(
        disclosure=disclosure,
        figures=[],
        export_path=tmp_path / "formula-export.docx",
        project_dir=tmp_path,
    )

    document = Document(str(output_path))
    formula_table = document.tables[0]
    assert len(formula_table.columns) == 3
    assert formula_table.cell(0, 0).text == ""
    assert formula_table.cell(0, 2).text == "(1)"
    formula_cell_pr = formula_table.cell(0, 1)._tc.tcPr
    assert formula_cell_pr is not None
    formula_cell_width = int(formula_cell_pr.tcW.w)
    assert formula_cell_width == int(docx_export.FORMULA_BODY_COL_IN * 1440)
    assert docx_export.FORMULA_MAX_WIDTH_IN < docx_export.FORMULA_BODY_COL_IN


def test_docx_export_renders_all_inline_math_as_images(tmp_path: Path) -> None:
    disclosure = build_initial_disclosure("inline math export")
    disclosure["sections"][0]["blocks"] = [
        {
            "id": "blk_inline_math_test",
            "type": "paragraph",
            "text": "condition $U > T$ and length $L_i$",
        }
    ]

    output_path = export_disclosure_docx(
        disclosure=disclosure,
        figures=[],
        export_path=tmp_path / "inline-math-export.docx",
        project_dir=tmp_path,
    )

    document = Document(str(output_path))
    paragraph = next(paragraph for paragraph in document.paragraphs if "condition" in paragraph.text)
    assert paragraph.text == "condition  and length "
    assert len(document.inline_shapes) == 2
    assert not any(run.font.name == "Cambria Math" for run in paragraph.runs)
    assert not list((tmp_path / "exports").glob("docx-assets-*"))


def test_docx_export_uses_existing_drawio_figure_png(tmp_path: Path) -> None:
    disclosure = build_initial_disclosure("figure export")
    disclosure["sections"][-1]["blocks"] = [
        {
            "id": "blk_figure_test",
            "type": "figure",
            "figure_id": "fig_000001",
        }
    ]
    render_path = tmp_path / "assets" / "figures" / "fig_000001" / "render.png"
    render_path.parent.mkdir(parents=True)
    render_path.write_bytes(PNG_BYTES)
    figures = [
        {
            "figure_id": "fig_000001",
            "label": "图1",
            "title": "系统结构示意图",
            "source": {
                "type": "drawio",
                "path": "assets/figures/fig_000001/diagram.drawio",
                "updated_at": "2026-07-09T10:00:00.000000+08:00",
            },
            "render": {
                "type": "png",
                "path": "assets/figures/fig_000001/render.png",
                "width": 1500,
                "height": 900,
            },
        }
    ]

    output_path = export_disclosure_docx(
        disclosure=disclosure,
        figures=figures,
        export_path=tmp_path / "figure-export.docx",
        project_dir=tmp_path,
    )

    document = Document(str(output_path))
    assert len(document.inline_shapes) == 1
    assert any(paragraph.text == "图1 系统结构示意图" for paragraph in document.paragraphs)
    assert not list((tmp_path / "exports").glob("docx-assets-*"))


def test_docx_export_rejects_missing_declared_figure_png(tmp_path: Path) -> None:
    disclosure = build_initial_disclosure("missing figure export")
    disclosure["sections"][-1]["blocks"] = [
        {
            "id": "blk_missing_figure",
            "type": "figure",
            "figure_id": "fig_000001",
        }
    ]
    figures = [
        {
            "figure_id": "fig_000001",
            "label": "图1",
            "title": "缺失截图",
            "render": {
                "type": "png",
                "path": "assets/figures/fig_000001/render.png",
                "width": 1500,
                "height": 900,
            },
        }
    ]

    with pytest.raises(DocxExportError, match="figure render file is missing: fig_000001"):
        export_disclosure_docx(
            disclosure=disclosure,
            figures=figures,
            export_path=tmp_path / "missing-figure-export.docx",
            project_dir=tmp_path,
        )


def test_docx_export_keeps_placeholder_for_unknown_figure(tmp_path: Path) -> None:
    disclosure = build_initial_disclosure("unknown figure export")
    disclosure["sections"][-1]["blocks"] = [
        {
            "id": "blk_unknown_figure",
            "type": "figure",
            "figure_id": "fig_000001",
        }
    ]

    output_path = export_disclosure_docx(
        disclosure=disclosure,
        figures=[],
        export_path=tmp_path / "unknown-figure-export.docx",
        project_dir=tmp_path,
    )

    document = Document(str(output_path))
    assert len(document.inline_shapes) == 0
    assert any(paragraph.text == "fig_000001" for paragraph in document.paragraphs)


def test_workspace_docx_export_uses_unlocked_stable_figure_snapshot(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    def render(_self: WorkspaceStore, *, input_path: Path, output_path: Path) -> dict:
        assert input_path.is_file()
        output_path.write_bytes(PNG_BYTES)
        return {"status": "success", "output": {"path": str(output_path)}}

    monkeypatch.setattr(WorkspaceStore, "_render_drawio_file", render)
    store = WorkspaceStore(tmp_path / "data", "Test User", "test@example.com")
    project = store.create_project("concurrent figure export")
    project_id = project.project_id
    drawio_xml = (Path(__file__).parent / "fixtures" / "figure-smoke.drawio").read_text(encoding="utf-8")
    created_result = store.create_figure(project_id, title="系统结构示意图", drawio_xml=drawio_xml)
    assert created_result["status"] == "success"
    created = created_result["output"]["figure"]

    disclosure = store.get_disclosure(project_id)
    disclosure["sections"][-1]["blocks"] = [
        {
            "id": "blk_concurrent_figure",
            "type": "figure",
            "figure_id": created["figure_id"],
        }
    ]
    store.save_disclosure(project_id, disclosure)
    old_render_path = store.figure_render_file(project_id, created["figure_id"])
    assert old_render_path.is_file()

    snapshot_ready = threading.Event()
    allow_docx_generation = threading.Event()
    captured_snapshot: dict[str, Path] = {}

    def delayed_export(**kwargs: object) -> Path:
        figures = kwargs["figures"]
        assert isinstance(figures, list)
        render_record = figures[0]["render"]
        snapshot_path = (Path(kwargs["project_dir"]) / render_record["path"]).resolve()
        captured_snapshot["path"] = snapshot_path
        assert snapshot_path != old_render_path.resolve()
        assert snapshot_path.is_file()
        snapshot_ready.set()
        assert allow_docx_generation.wait(timeout=5)
        assert not old_render_path.exists()
        assert snapshot_path.read_bytes() == PNG_BYTES
        return export_disclosure_docx(**kwargs)

    monkeypatch.setattr(workspace_store_module, "export_disclosure_docx", delayed_export)

    with ThreadPoolExecutor(max_workers=2) as pool:
        export_future = pool.submit(store.export_docx, project_id)
        assert snapshot_ready.wait(timeout=5)
        write_future = pool.submit(
            store.write_figure,
            project_id,
            created["figure_id"],
            title="更新后的系统结构示意图",
            drawio_xml=drawio_xml.replace('value="Source"', 'value="Updated Source"'),
            expected_drawio_updated_at=created["source"]["updated_at"],
        )
        try:
            write_result = write_future.result(timeout=5)
        finally:
            allow_docx_generation.set()
        output_path = export_future.result(timeout=5)

    assert write_result["status"] == "success"
    assert not old_render_path.exists()
    assert not captured_snapshot["path"].exists()
    document = Document(str(output_path))
    assert len(document.inline_shapes) == 1


def test_docx_export_cleans_asset_directory_when_renderer_fails(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    disclosure = build_initial_disclosure("asset failure export")
    disclosure["sections"][0]["blocks"] = [
        {
            "id": "blk_inline_math_test",
            "type": "paragraph",
            "text": "condition $U > T$",
        }
    ]

    def fail_renderer(*_args: object, **_kwargs: object) -> subprocess.CompletedProcess[str]:
        return subprocess.CompletedProcess(args="", returncode=1, stdout="", stderr="renderer failed")

    monkeypatch.setattr(docx_export.subprocess, "run", fail_renderer)

    with pytest.raises(DocxExportError, match="renderer failed"):
        export_disclosure_docx(
            disclosure=disclosure,
            figures=[],
            export_path=tmp_path / "failed-export.docx",
            project_dir=tmp_path,
        )

    assert not list((tmp_path / "exports").glob("docx-assets-*"))


def test_docx_export_rejects_renderer_manifest_missing_assets(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    disclosure = build_initial_disclosure("missing asset export")
    disclosure["sections"][0]["blocks"] = [
        {
            "id": "blk_inline_math_test",
            "type": "paragraph",
            "text": "condition $U > T$",
        }
    ]

    def write_incomplete_manifest(args: list[str], **_kwargs: object) -> subprocess.CompletedProcess[str]:
        manifest_path = Path(args[args.index("--manifest") + 1])
        manifest_path.write_text('{"assets": {}}\n', encoding="utf-8")
        return subprocess.CompletedProcess(args=args, returncode=0, stdout="", stderr="")

    monkeypatch.setattr(docx_export.subprocess, "run", write_incomplete_manifest)

    with pytest.raises(DocxExportError, match="did not return asset"):
        export_disclosure_docx(
            disclosure=disclosure,
            figures=[],
            export_path=tmp_path / "missing-asset-export.docx",
            project_dir=tmp_path,
        )

    assert not list((tmp_path / "exports").glob("docx-assets-*"))


def test_docx_export_rejects_renderer_manifest_missing_files(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    disclosure = build_initial_disclosure("missing file export")
    disclosure["sections"][0]["blocks"] = [
        {
            "id": "blk_inline_math_test",
            "type": "paragraph",
            "text": "condition $U > T$",
        }
    ]

    def write_missing_file_manifest(args: list[str], **_kwargs: object) -> subprocess.CompletedProcess[str]:
        input_path = Path(args[args.index("--input") + 1])
        output_dir = Path(args[args.index("--output") + 1])
        manifest_path = Path(args[args.index("--manifest") + 1])
        payload = json.loads(input_path.read_text(encoding="utf-8"))
        asset_id = payload["items"][0]["id"]
        manifest_path.write_text(
            json.dumps(
                {
                    "assets": {
                        asset_id: {
                            "id": asset_id,
                            "kind": "inline_formula",
                            "path": str(output_dir / "does-not-exist.png"),
                            "width_px": 20,
                            "height_px": 10,
                        }
                    }
                }
            )
            + "\n",
            encoding="utf-8",
        )
        return subprocess.CompletedProcess(args=args, returncode=0, stdout="", stderr="")

    monkeypatch.setattr(docx_export.subprocess, "run", write_missing_file_manifest)

    with pytest.raises(DocxExportError, match="missing file"):
        export_disclosure_docx(
            disclosure=disclosure,
            figures=[],
            export_path=tmp_path / "missing-file-export.docx",
            project_dir=tmp_path,
        )

    assert not list((tmp_path / "exports").glob("docx-assets-*"))


def test_docx_export_rejects_renderer_manifest_paths_outside_output_dir(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    disclosure = build_initial_disclosure("outside asset export")
    disclosure["sections"][0]["blocks"] = [
        {
            "id": "blk_inline_math_test",
            "type": "paragraph",
            "text": "condition $U > T$",
        }
    ]
    outside_asset = tmp_path / "outside.png"
    outside_asset.write_bytes(b"not a real png")

    def write_outside_file_manifest(args: list[str], **_kwargs: object) -> subprocess.CompletedProcess[str]:
        input_path = Path(args[args.index("--input") + 1])
        manifest_path = Path(args[args.index("--manifest") + 1])
        payload = json.loads(input_path.read_text(encoding="utf-8"))
        asset_id = payload["items"][0]["id"]
        manifest_path.write_text(
            json.dumps(
                {
                    "assets": {
                        asset_id: {
                            "id": asset_id,
                            "kind": "inline_formula",
                            "path": str(outside_asset),
                            "width_px": 20,
                            "height_px": 10,
                        }
                    }
                }
            )
            + "\n",
            encoding="utf-8",
        )
        return subprocess.CompletedProcess(args=args, returncode=0, stdout="", stderr="")

    monkeypatch.setattr(docx_export.subprocess, "run", write_outside_file_manifest)

    with pytest.raises(DocxExportError, match="out-of-directory path"):
        export_disclosure_docx(
            disclosure=disclosure,
            figures=[],
            export_path=tmp_path / "outside-asset-export.docx",
            project_dir=tmp_path,
        )

    assert not list((tmp_path / "exports").glob("docx-assets-*"))


def test_docx_export_cleans_asset_directory_when_document_build_fails(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    disclosure = build_initial_disclosure("document failure export")
    disclosure["sections"][0]["blocks"] = [
        {
            "id": "blk_inline_math_test",
            "type": "paragraph",
            "text": "condition $U > T$",
        }
    ]

    def fail_render_section(*args: object, **kwargs: object) -> None:
        raise RuntimeError("document build failed")

    monkeypatch.setattr(docx_export, "render_section", fail_render_section)

    with pytest.raises(RuntimeError, match="document build failed"):
        export_disclosure_docx(
            disclosure=disclosure,
            figures=[],
            export_path=tmp_path / "failed-document-export.docx",
            project_dir=tmp_path,
        )

    assert not list((tmp_path / "exports").glob("docx-assets-*"))


def test_docx_export_keeps_tables_editable(tmp_path: Path) -> None:
    disclosure = build_initial_disclosure("table export")
    disclosure["sections"][0]["blocks"].append(
        {
            "id": "blk_900001",
            "type": "table",
            "columns": ["模块", "作用"],
            "rows": [["前端交互", "展示预览"], ["文档领域", "维护结构化正文"]],
        }
    )

    output_path = export_disclosure_docx(
        disclosure=disclosure,
        figures=[],
        export_path=tmp_path / "table-export.docx",
        project_dir=tmp_path,
    )

    document = Document(str(output_path))
    assert len(document.tables) == 1
    table = document.tables[0]
    assert table.cell(0, 0).text == "模块"
    assert table.cell(0, 1).text == "作用"
    assert table.cell(1, 0).text == "前端交互"
    assert table.cell(2, 1).text == "维护结构化正文"
